import sys
import traceback

# Check for required libraries
try:
    import pandas as pd
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from PyQt5.QtWidgets import (QApplication, QMainWindow, QVBoxLayout, QHBoxLayout, 
                                 QWidget, QPushButton, QTableWidget, QTableWidgetItem, 
                                 QFileDialog, QMessageBox, QLineEdit, QLabel, QComboBox,
                                 QHeaderView,  QMenu, QAction, QDialog, QTextEdit,
                                 QSplitter, QGroupBox, QCheckBox, QShortcut)
    from PyQt5.QtCore import Qt, QSettings, QTimer
    from PyQt5.QtGui import QKeySequence, QColor, QBrush
    from PyQt5 import QtGui
    from SplashScreen_   import SplashScreen
    from DataProcessor_  import DataProcessor
    from Settings_       import SettingsDialog
    from Statistics_     import StatisticsDialog
    from SheetSelection_ import SheetSelectionDialog
    from NewFile_        import NewFileDialog
    from AddColumn_      import AddColumnDialog
    from AddRow_         import AddRowDialog
    from ColumnSelection_ import ColumnSelectionDialog
    from Theme_           import apply_custom_theme, apply_dark_theme
    from ColumnFormatting_ import ColumnFormattingDialog
    from FormatApplier_ import FormatApplier
    from AdvancedFormatting_ import AdvancedFormattingDialog
    from DataTransformation_ import DataTransformationDialog
    from Help_           import HelpDialog
    from AutoSave_       import AutoSaveManager, VersionHistoryDialog, AutoSaveSettingsDialog
    from UndoRedo_       import UndoRedoManager, UndoRedoHistoryDialog
    from QuickActions_   import QuickActionsMenu
    from Favorites_      import FavoritesManager, FavoritesDialog
    from ColumnFreeze_   import ColumnFreezeManager, ColumnFreezeDialog
    from SplitView_      import SplitViewManager, SplitViewOptionsDialog
    from AIDialog_       import AIFeaturesDialog
    from AIDialog_       import AIFeaturesDialog
    from CloudSync_      import CloudSyncDialog
    from Visualization_  import VisualizationManager, ChartDialog, DashboardDialog
    from DataValidation_ import ValidationManager, ValidationDialog, DropdownDelegate, ValidationRule
    import os
    print("All PyQt5 modules imported successfully")
except ImportError as e:
    print(f"Import error: {e}")
    print("Please install required packages:")
    print("pip install PyQt5 pandas openpyxl")
    input("Press Enter to exit...")
    sys.exit(1)

# Try to import openpyxl
try:
    from openpyxl import Workbook, load_workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    OPENPYXL_AVAILABLE = True
    print("openpyxl imported successfully")
except ImportError:
    print("openpyxl not available - Excel advanced features disabled")
    OPENPYXL_AVAILABLE = False

try:
    import numpy as np
    NUMPY_AVAILABLE = True
except ImportError:
    NUMPY_AVAILABLE = False


class ExcelEditor(QMainWindow):
    def __init__(self):
        super().__init__()
        print("Initializing ExcelEditor...")
        try:
            self.df = None
            self.filtered_df = None
            self.current_file_path = None
            self.is_modified = False
            
            # Initialize productivity managers
            self.auto_save_manager = AutoSaveManager(self)
            self.undo_redo_manager = UndoRedoManager(self)
            self.favorites_manager = FavoritesManager(self)
            self.freeze_manager = ColumnFreezeManager(self)
            self.favorites_manager = FavoritesManager(self)
            self.freeze_manager = ColumnFreezeManager(self)
            self.split_view_manager = SplitViewManager(self)
            self.validation_manager = ValidationManager()
            
            self.format_settings = {}  # Store column formatting settings
            
            self.MAX_DISPLAY_ROWS = 1000
            self.is_updating = False
            
            # Load settings
            self.settings = QSettings("DataTools", "ExcelEditor")
            self.bg_color = self.settings.value("theme/bg_color", "#313B2F")
            self.accent_color = self.settings.value("theme/accent_color", "#FBA002")
            self.auto_save_enabled = self.settings.value("auto_save/enabled", True, type=bool)
            self.auto_save_interval = int(self.settings.value("auto_save/interval", 5)) # minutes
            
            self.auto_save_interval = int(self.settings.value("auto_save/interval", 5)) # minutes

            # Recent files list
            self.recent_files = self.settings.value("recent_files", [])
            if not isinstance(self.recent_files, list):
                self.recent_files = []
            
            print("Setting up UI...")
            self.init_ui()
            print("Creating menu bar...")
            self.create_menu_bar()
            #print("Creating toolbar...")
            #self.create_toolbar()
            print("Creating status bar...")
            self.create_status_bar()
            print("Setting up keyboard shortcuts...")
            self.setup_keyboard_shortcuts()
            
            # Initialize Quick Actions Context Menu
            self.quick_actions = QuickActionsMenu(self)
            

            
            print("ExcelEditor initialized successfully")
        except Exception as e:
            print(f"Error initializing ExcelEditor: {e}")
            traceback.print_exc()
            raise
        
    def init_ui(self):
        self.setWindowTitle("Excel Editor  Pro - New File")
        self.setGeometry(100, 100, 1200, 800)
        
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        main_layout = QVBoxLayout()
        central_widget.setLayout(main_layout)
        
        control_panel = self.create_control_panel()
        main_layout.addWidget(control_panel)
        
        splitter = QSplitter(Qt.Horizontal)
        
        self.table_widget = QTableWidget()
        self.table_widget.setSortingEnabled(False)
        self.table_widget.setAlternatingRowColors(True)
        self.table_widget.setSelectionBehavior(QTableWidget.SelectRows)
        #self.table_widget.setDragEnabled(True)
        #self.table_widget.setAcceptDrops(True)
        #self.table_widget.setViewportExternalDragDrop(True)
        #self.table_widget.setDragDropMode(QTableWidget.InternalMove)
        #self.table_widget.model().rowsMoved.connect(self.on_rows_moved)
        self.table_widget.cellChanged.connect(self.on_cell_changed)
        self.table_widget.setVerticalScrollMode(QTableWidget.ScrollPerPixel)
        self.table_widget.setHorizontalScrollMode(QTableWidget.ScrollPerPixel)
        
        splitter.addWidget(self.table_widget)
        
        info_panel = self.create_info_panel()
        splitter.addWidget(info_panel)
        
        splitter.setSizes([800, 400])
        main_layout.addWidget(splitter)
        
    #def create_toolbar(self):
        #try:
            #toolbar = QToolBar("Main Toolbar")
            #self.addToolBar(toolbar)
            
            #new_action = toolbar.addAction("New")
            #new_action.triggered.connect(self.new_file)
            #new_action.setShortcut(QKeySequence.New)
            
            #open_action = toolbar.addAction("Open")
            #open_action.triggered.connect(self.load_file)
            #open_action.setShortcut(QKeySequence.Open)
            
            #save_action = toolbar.addAction("Save")
            #save_action.triggered.connect(self.save_file)
            #save_action.setShortcut(QKeySequence.Save)
            
            #toolbar.addSeparator()
            
            #toolbar.addAction("Add Row", self.add_row)
            #toolbar.addAction("Add Column", self.add_column)
            #toolbar.addAction("Delete Row", self.delete_row)
        #except Exception as e:
            #print(f"Error creating toolbar: {e}")
    
    def create_status_bar(self):
        try:
            status_bar = self.statusBar()
            self.cell_info_label = QLabel("Ready")
            status_bar.addWidget(self.cell_info_label)
            status_bar.showMessage("Ready - Create a new file or load existing one")
        except Exception as e:
            print(f"Error creating status bar: {e}")
        
    def create_control_panel(self):
        panel = QGroupBox("Controls")
        layout = QVBoxLayout()
        layout.setSpacing(8)
        layout.setContentsMargins(10, 15, 10, 10)
        
        file_row = QHBoxLayout()
        file_row.setSpacing(8)
        
        self.new_btn = QPushButton("New File")
        self.new_btn.clicked.connect(self.new_file)
        self.new_btn.setMaximumHeight(32)
        self.new_btn.setToolTip("Create a new Excel or CSV file (Ctrl+N)")
        file_row.addWidget(self.new_btn)
        
        self.load_btn = QPushButton("Open File")
        self.load_btn.clicked.connect(self.load_file)
        self.load_btn.setMaximumHeight(32)
        self.load_btn.setToolTip("Open an existing Excel or CSV file (Ctrl+O)")
        file_row.addWidget(self.load_btn)
        
        self.save_btn = QPushButton("Save")
        self.save_btn.clicked.connect(self.save_file)
        self.save_btn.setEnabled(False)
        self.save_btn.setMaximumHeight(32)
        self.save_btn.setToolTip("Save current file (Ctrl+S)")
        file_row.addWidget(self.save_btn)
        
        self.save_as_btn = QPushButton("Save As...")
        self.save_as_btn.clicked.connect(self.save_as_file)
        self.save_as_btn.setEnabled(False)
        self.save_as_btn.setMaximumHeight(32)
        self.save_as_btn.setToolTip("Save with a new name or format (Ctrl+Shift+S)")
        file_row.addWidget(self.save_as_btn)
        
        #self.settings_btn = QPushButton("Settings")
        #self.settings_btn.clicked.connect(self.open_settings)
        #self.settings_btn.setMaximumHeight(32)
        #file_row.addWidget(self.settings_btn)
        
        file_row.addStretch()
        layout.addLayout(file_row)
        
        edit_row = QHBoxLayout()
        edit_row.setSpacing(8)
        
        self.add_row_btn = QPushButton("Add Row")
        self.add_row_btn.clicked.connect(self.add_row)
        self.add_row_btn.setEnabled(False)
        self.add_row_btn.setMaximumHeight(32)
        self.add_row_btn.setToolTip("Add a new row to the spreadsheet (Ctrl+R)")
        edit_row.addWidget(self.add_row_btn)
        
        self.add_column_btn = QPushButton("Add Column")
        self.add_column_btn.clicked.connect(self.add_column)
        self.add_column_btn.setEnabled(False)
        self.add_column_btn.setMaximumHeight(32)
        self.add_column_btn.setToolTip("Add a new column to the spreadsheet (Ctrl+Shift+C)")
        edit_row.addWidget(self.add_column_btn)
        
        self.format_column_btn = QPushButton("Format Columns")
        self.format_column_btn.clicked.connect(self.format_columns)
        self.format_column_btn.setEnabled(False)
        self.format_column_btn.setMaximumHeight(32)
        self.format_column_btn.setToolTip("Format columns with fonts, colors, alignment, and number formats (Ctrl+Shift+F)")
        edit_row.addWidget(self.format_column_btn)
        
        self.advanced_format_btn = QPushButton("📈 Advanced Format")
        self.advanced_format_btn.clicked.connect(self.open_advanced_formatting)
        self.advanced_format_btn.setEnabled(False)
        self.advanced_format_btn.setMaximumHeight(32)
        self.advanced_format_btn.setToolTip("Access advanced formatting features (Ctrl+Alt+F)")
        edit_row.addWidget(self.advanced_format_btn)
        
        self.transform_btn = QPushButton("🔧 Transform Data")
        self.transform_btn.clicked.connect(self.open_data_transformation)
        self.transform_btn.setEnabled(False)
        self.transform_btn.setMaximumHeight(32)
        self.transform_btn.setToolTip("Transform, clean, and manipulate data (Ctrl+T)")
        edit_row.addWidget(self.transform_btn)

        self.visualize_btn = QPushButton("📊 Visualize")
        self.visualize_btn.clicked.connect(self.open_visualization_menu)
        self.visualize_btn.setEnabled(False)
        self.visualize_btn.setMaximumHeight(32)
        self.visualize_btn.setToolTip("Create charts and visualizations")
        self.visualize_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0, stop:0 #E91E63, stop:1 #C2185B);
                color: white;
                font-weight: bold;
            }
            QPushButton:hover {
                 background: qlineargradient(x1:0, y1:0, x2:1, y2:0, stop:0 #F06292, stop:1 #E91E63);
            }
        """)
        edit_row.addWidget(self.visualize_btn)
        
        self.ai_btn = QPushButton("🤖 AI Assistant")
        self.ai_btn.clicked.connect(self.open_ai_features)
        self.ai_btn.setEnabled(False)
        self.ai_btn.setMaximumHeight(32)
        self.ai_btn.setToolTip("AI-powered data analysis and operations (Ctrl+Shift+A)")
        self.ai_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #6B46C1, stop:1 #9333EA);
                color: white;
                font-weight: bold;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #7C3AED, stop:1 #A855F7);
            }
        """)
        edit_row.addWidget(self.ai_btn)
        
        self.cloud_btn = QPushButton("☁️ Cloud Sync")
        self.cloud_btn.clicked.connect(self.open_cloud_sync)
        self.cloud_btn.setMaximumHeight(32)
        self.cloud_btn.setToolTip("Upload, download, and sync files with cloud storage (Ctrl+Shift+U)")
        self.cloud_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #0EA5E9, stop:1 #06B6D4);
                color: white;
                font-weight: bold;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #0284C7, stop:1 #0891B2);
            }
        """)
        edit_row.addWidget(self.cloud_btn)
        
        self.delete_row_btn = QPushButton("Delete Row")
        self.delete_row_btn.clicked.connect(self.delete_row)
        self.delete_row_btn.setEnabled(False)
        self.delete_row_btn.setMaximumHeight(32)
        self.delete_row_btn.setToolTip("Delete the selected row (Delete key)")
        edit_row.addWidget(self.delete_row_btn)
        
        self.move_up_btn = QPushButton("Move Up ▲")
        self.move_up_btn.clicked.connect(self.move_row_up)
        self.move_up_btn.setEnabled(False)
        self.move_up_btn.setMaximumHeight(32)
        self.move_up_btn.setToolTip("Move selected row up one position")
        edit_row.addWidget(self.move_up_btn)

        self.move_down_btn = QPushButton("Move Down ▼")
        self.move_down_btn.clicked.connect(self.move_row_down)
        self.move_down_btn.setEnabled(False)
        self.move_down_btn.setMaximumHeight(32)
        self.move_down_btn.setToolTip("Move selected row down one position")
        edit_row.addWidget(self.move_down_btn)
        
        edit_row.addStretch()
        layout.addLayout(edit_row)
        
        filter_row = QHBoxLayout()
        filter_row.setSpacing(8)
        
        filter_row.addWidget(QLabel("Filter Column:"))
        self.filter_column_combo = QComboBox()
        self.filter_column_combo.setMinimumWidth(150)
        self.filter_column_combo.setMaximumHeight(28)
        self.filter_column_combo.setToolTip("Select which column to filter by")
        filter_row.addWidget(self.filter_column_combo)
        
        filter_row.addWidget(QLabel("Filter Value:"))
        self.filter_input = QLineEdit()
        self.filter_input.setPlaceholderText("Enter filter value...")
        self.filter_input.returnPressed.connect(self.apply_filter)
        self.filter_input.setMaximumHeight(28)
        self.filter_input.setToolTip("Enter text to search for in the selected column (Ctrl+F to focus)")
        filter_row.addWidget(self.filter_input)
        
        self.apply_filter_btn = QPushButton("Apply Filter")
        self.apply_filter_btn.clicked.connect(self.apply_filter)
        self.apply_filter_btn.setMaximumHeight(28)
        self.apply_filter_btn.setToolTip("Show only rows matching the filter criteria")
        filter_row.addWidget(self.apply_filter_btn)
        
        self.clear_filter_btn = QPushButton("Clear Filter")
        self.clear_filter_btn.clicked.connect(self.clear_filter)
        self.clear_filter_btn.setEnabled(False)
        self.clear_filter_btn.setMaximumHeight(28)
        self.clear_filter_btn.setToolTip("Remove filter and show all rows")
        filter_row.addWidget(self.clear_filter_btn)
        
        filter_row.addStretch()
        layout.addLayout(filter_row)
        
        panel.setLayout(layout)
        panel.setMaximumHeight(120)
        return panel
    
    def create_info_panel(self):
        panel = QGroupBox("Data Information")
        layout = QVBoxLayout()
        
        self.info_text = QTextEdit()
        self.info_text.setReadOnly(True)
        self.info_text.setMaximumHeight(200)
        layout.addWidget(self.info_text)
        
        self.stats_btn = QPushButton("Show Detailed Statistics")
        self.stats_btn.clicked.connect(self.show_statistics)
        self.stats_btn.setEnabled(False)
        self.stats_btn.setToolTip("View comprehensive statistics and analysis of your data (Ctrl+I)")
        layout.addWidget(self.stats_btn)
        
        export_group = QGroupBox("Export Options")
        export_layout = QVBoxLayout()
        
        self.include_index_cb = QCheckBox("Include Row Index")
        self.include_index_cb.setChecked(False)
        self.include_index_cb.setToolTip("Include row numbers when saving to Excel/CSV")
        export_layout.addWidget(self.include_index_cb)
        
        export_group.setLayout(export_layout)
        layout.addWidget(export_group)
        
        layout.addStretch()
        panel.setLayout(layout)
        return panel
    
    def create_menu_bar(self):
        try:
            menubar = self.menuBar()
            
            file_menu = menubar.addMenu('File')
            
            new_action = QAction('New', self)
            new_action.setShortcut('Ctrl+N')
            new_action.triggered.connect(self.new_file)
            file_menu.addAction(new_action)
            
            open_action = QAction('Open', self)
            open_action.setShortcut('Ctrl+O')
            open_action.triggered.connect(self.load_file)
            file_menu.addAction(open_action)
            
            # Recent Files Submenu
            self.recent_menu = QMenu('Open Recent', self)
            file_menu.addMenu(self.recent_menu)
            self.update_recent_files_menu()
            
            file_menu.addSeparator()
            
            save_action = QAction('Save', self)
            save_action.setShortcut('Ctrl+S')
            save_action.triggered.connect(self.save_file)
            file_menu.addAction(save_action)
            
            save_as_action = QAction('Save As...', self)
            save_as_action.setShortcut('Ctrl+Shift+S')
            save_as_action.triggered.connect(self.save_as_file)
            file_menu.addAction(save_as_action)
            
            export_doc_action = QAction('Export Report (DOC)', self)
            export_doc_action.triggered.connect(self.export_to_doc)
            file_menu.addAction(export_doc_action)
            
            file_menu.addSeparator()
            
            # Favorites submenu
            favorites_menu = file_menu.addMenu('Favorites')
            
            view_favorites_action = QAction('View Favorites', self)
            view_favorites_action.setShortcut('Ctrl+B')
            view_favorites_action.triggered.connect(self.open_favorites)
            favorites_menu.addAction(view_favorites_action)
            
            add_favorite_action = QAction('Add Current File to Favorites', self)
            add_favorite_action.setShortcut('Ctrl+D')
            add_favorite_action.triggered.connect(self.add_current_to_favorites)
            favorites_menu.addAction(add_favorite_action)
            
            file_menu.addSeparator()
            
            # Version History
            version_history_action = QAction('Version History', self)
            version_history_action.setShortcut('Ctrl+H')
            version_history_action.triggered.connect(self.open_version_history)
            file_menu.addAction(version_history_action)
            
            file_menu.addSeparator()
            
            exit_action = QAction('Exit', self)
            exit_action.setShortcut('Ctrl+Q')
            exit_action.triggered.connect(self.close)
            file_menu.addAction(exit_action)
            
            edit_menu = menubar.addMenu('Edit')
            
            # Undo/Redo
            undo_action = QAction('↶ Undo', self)
            undo_action.setShortcut('Ctrl+Z')
            undo_action.triggered.connect(self.undo)
            edit_menu.addAction(undo_action)
            
            redo_action = QAction('↷ Redo', self)
            redo_action.setShortcut('Ctrl+Y')
            redo_action.triggered.connect(self.redo)
            edit_menu.addAction(redo_action)
            
            history_action = QAction('Edit History', self)
            history_action.setShortcut('Ctrl+Shift+H')
            history_action.triggered.connect(self.open_edit_history)
            edit_menu.addAction(history_action)
            
            edit_menu.addSeparator()
            
            add_row_action = QAction('Add Row', self)
            add_row_action.setShortcut('Ctrl+R')
            add_row_action.triggered.connect(self.add_row)
            edit_menu.addAction(add_row_action)
            
            add_column_action = QAction('Add Column', self)
            add_column_action.setShortcut('Ctrl+Shift+C')
            add_column_action.triggered.connect(self.add_column)
            edit_menu.addAction(add_column_action)
            
            format_column_action = QAction('Format Columns', self)
            format_column_action.setShortcut('Ctrl+Shift+F')
            format_column_action.triggered.connect(self.format_columns)
            edit_menu.addAction(format_column_action)
            
            advanced_format_action = QAction('Advanced Formatting', self)
            advanced_format_action.setShortcut('Ctrl+Alt+F')
            advanced_format_action.triggered.connect(self.open_advanced_formatting)
            edit_menu.addAction(advanced_format_action)
            
            validation_action = QAction('Data Validation', self)
            validation_action.setShortcut('Ctrl+Shift+D')
            validation_action.triggered.connect(self.open_validation_dialog)
            edit_menu.addAction(validation_action)
            
            edit_menu.addSeparator()
            
            transform_action = QAction('Data Transformation', self)
            transform_action.setShortcut('Ctrl+T')
            transform_action.triggered.connect(self.open_data_transformation)
            edit_menu.addAction(transform_action)
            
            delete_row_action = QAction('Delete Row', self)
            delete_row_action.setShortcut('Delete')
            delete_row_action.triggered.connect(self.delete_row)
            edit_menu.addAction(delete_row_action)

            # Visualization Menu
            viz_menu = menubar.addMenu('Visualize')
            
            chart_action = QAction('Create Chart', self)
            chart_action.setShortcut('Ctrl+Shift+V')
            chart_action.triggered.connect(self.create_chart)
            viz_menu.addAction(chart_action)
            
            dashboard_action = QAction('Dashboard View', self)
            dashboard_action.setShortcut('Ctrl+Alt+V')
            dashboard_action.triggered.connect(self.open_dashboard)
            viz_menu.addAction(dashboard_action)

            # AI Menu
            ai_menu = menubar.addMenu('AI')
            
            ai_features_action = QAction('AI Data Assistant', self)
            ai_features_action.setShortcut('Ctrl+Shift+A')
            ai_features_action.triggered.connect(self.open_ai_features)
            ai_menu.addAction(ai_features_action)
            
            ai_menu.addSeparator()
            
            ai_analyze_action = QAction('Quick Analysis', self)
            ai_analyze_action.setShortcut('Ctrl+Alt+A')
            ai_analyze_action.triggered.connect(self.quick_ai_analysis)
            ai_menu.addAction(ai_analyze_action)
            
            ai_clean_action = QAction('Clean Data', self)
            ai_clean_action.triggered.connect(self.smart_clean_data)
            ai_menu.addAction(ai_clean_action)
            
            ai_suggest_action = QAction('Get Suggestions', self)
            ai_suggest_action.triggered.connect(self.get_ai_suggestions)
            ai_menu.addAction(ai_suggest_action)
            
            # Cloud menu
            cloud_menu = menubar.addMenu('Cloud')
            
            cloud_sync_action = QAction('Cloud Sync', self)
            cloud_sync_action.setShortcut('Ctrl+Shift+U')
            cloud_sync_action.triggered.connect(self.open_cloud_sync)
            cloud_menu.addAction(cloud_sync_action)
            
            cloud_menu.addSeparator()
            
            quick_upload_action = QAction('Quick Upload', self)
            quick_upload_action.setShortcut('Ctrl+U')
            quick_upload_action.triggered.connect(self.quick_upload_to_cloud)
            cloud_menu.addAction(quick_upload_action)
            
            download_from_cloud_action = QAction('Download from Cloud', self)
            download_from_cloud_action.triggered.connect(self.download_from_cloud)
            cloud_menu.addAction(download_from_cloud_action)
            
            cloud_menu.addSeparator()
            
            auto_backup_action = QAction('Auto-Backup Settings', self)
            auto_backup_action.triggered.connect(self.configure_cloud_backup)
            cloud_menu.addAction(auto_backup_action)
            
            view_menu = menubar.addMenu('View')
            
            # Split View submenu
            split_view_menu = view_menu.addMenu('Split View')
            
            split_horizontal_action = QAction('Split Horizontal', self)
            split_horizontal_action.triggered.connect(lambda: self.split_view_manager.create_split_view('horizontal'))
            split_view_menu.addAction(split_horizontal_action)
            
            split_vertical_action = QAction('Split Vertical', self)
            split_vertical_action.triggered.connect(lambda: self.split_view_manager.create_split_view('vertical'))
            split_view_menu.addAction(split_vertical_action)
            
            split_view_menu.addSeparator()
            
            close_split_action = QAction('Close Split View', self)
            close_split_action.triggered.connect(self.split_view_manager.close_split_view)
            split_view_menu.addAction(close_split_action)
            
            view_menu.addSeparator()
            
            # Freeze Panes
            freeze_panes_action = QAction('Freeze Panes', self)
            freeze_panes_action.setShortcut('Ctrl+Shift+P')
            freeze_panes_action.triggered.connect(self.open_freeze_panes)
            view_menu.addAction(freeze_panes_action)
            
            view_menu.addSeparator()
            
            stats_action = QAction('Show Statistics', self)
            stats_action.setShortcut('Ctrl+I')
            stats_action.triggered.connect(self.show_statistics)
            view_menu.addAction(stats_action)
            
            view_menu.addSeparator()
            
            refresh_action = QAction('Refresh View', self)
            refresh_action.setShortcut('F5')
            refresh_action.triggered.connect(self.refresh_view)
            view_menu.addAction(refresh_action)
            
            # Settings menu
            self.settings_menu = menubar.addMenu('&Settings')
            
            settings_action = QAction('&General Settings', self)
            settings_action.setShortcut('Ctrl+,')
            settings_action.triggered.connect(self.open_settings)
            self.settings_menu.addAction(settings_action)
            
            autosave_settings_action = QAction('Auto-Save Settings', self)
            autosave_settings_action.triggered.connect(self.open_autosave_settings)
            self.settings_menu.addAction(autosave_settings_action)
            
            # Help menu
            help_menu = menubar.addMenu('&Help')
            
            help_action = QAction('Help & Documentation', self)
            help_action.setShortcut('F1')
            help_action.triggered.connect(self.show_help)
            help_menu.addAction(help_action)
            
            keyboard_shortcuts_action = QAction('Keyboard Shortcuts', self)
            keyboard_shortcuts_action.setShortcut('Ctrl+/')
            keyboard_shortcuts_action.triggered.connect(self.show_keyboard_shortcuts)
            help_menu.addAction(keyboard_shortcuts_action)
            
            help_menu.addSeparator()
            
            about_action = QAction('About Excel Editor Pro', self)
            about_action.triggered.connect(self.show_about)
            help_menu.addAction(about_action)
        except Exception as e:
            print(f"Error creating menu bar: {e}")
    
    def update_recent_files_menu(self):
        """Update the 'Open Recent' submenu with the latest files."""
        try:
            self.recent_menu.clear()
            
            if not self.recent_files:
                no_recent_action = QAction("No Recent Files", self)
                no_recent_action.setEnabled(False)
                self.recent_menu.addAction(no_recent_action)
                return

            for file_path in self.recent_files:
                if os.path.exists(file_path):
                    action = QAction(os.path.basename(file_path), self)
                    action.setData(file_path)
                    action.triggered.connect(lambda checked, path=file_path: self.open_recent_file(path))
                    self.recent_menu.addAction(action)
            
            if self.recent_menu.actions():
                self.recent_menu.addSeparator()
                clear_action = QAction("Clear Recent Files", self)
                clear_action.triggered.connect(self.clear_recent_files)
                self.recent_menu.addAction(clear_action)
            else:
                no_recent_action = QAction("No Recent Files", self)
                no_recent_action.setEnabled(False)
                self.recent_menu.addAction(no_recent_action)
        except Exception as e:
            print(f"Error updating recent files menu: {e}")

    def add_to_recent_files(self, file_path):
        """Add a file path to the recent files list and update settings."""
        try:
            if not file_path:
                return
                
            # Normalize path
            file_path = os.path.abspath(file_path)
            
            # Remove if already in list to move to top
            if file_path in self.recent_files:
                self.recent_files.remove(file_path)
            
            # Add to top
            self.recent_files.insert(0, file_path)
            
            # Keep only top 10
            self.recent_files = self.recent_files[:10]
            
            # Save to settings
            self.settings.setValue("recent_files", self.recent_files)
            
            # Update menu
            self.update_recent_files_menu()
        except Exception as e:
            print(f"Error adding to recent files: {e}")

    def open_recent_file(self, file_path):
        """Open a file from the recent files list."""
        try:
            if self.is_modified:
                reply = QMessageBox.question(
                    self, 'Unsaved Changes', 
                    'You have unsaved changes. Do you want to save them first?',
                    QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel,
                    QMessageBox.Cancel
                )
                
                if reply == QMessageBox.Cancel:
                    return
                elif reply == QMessageBox.Yes:
                    if not self.save_file():
                        return
            
            if os.path.exists(file_path):
                self.statusBar().showMessage(f"Loading recent file: {os.path.basename(file_path)}...")
                self.current_file_path = file_path
                self.start_file_loading(file_path)
            else:
                QMessageBox.warning(self, "File Not Found", f"The file '{file_path}' no longer exists.")
                # Remove from list if it doesn't exist
                if file_path in self.recent_files:
                    self.recent_files.remove(file_path)
                    self.settings.setValue("recent_files", self.recent_files)
                    self.update_recent_files_menu()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error opening recent file: {str(e)}")

    def clear_recent_files(self):
        """Clear the recent files list."""
        self.recent_files = []
        self.settings.setValue("recent_files", self.recent_files)
        self.update_recent_files_menu()

    def new_file(self):
        try:
            if self.is_modified:
                reply = QMessageBox.question(
                    self, 'Unsaved Changes', 
                    'You have unsaved changes. Do you want to save them first?',
                    QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel,
                    QMessageBox.Cancel
                )
                
                if reply == QMessageBox.Cancel:
                    return
                elif reply == QMessageBox.Yes:
                    if not self.save_file():
                        return
            
            dialog = NewFileDialog(self)
            if dialog.exec_() == QDialog.Accepted:
                settings = dialog.get_settings()
                self.create_new_file(settings)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error creating new file: {str(e)}")
    
    def create_new_file(self, settings):
        try:
            rows = settings['rows']
            cols = settings['columns']
            include_headers = settings['include_headers']
            
            if include_headers:
                columns = [f"Column {i+1}" for i in range(cols)]
            else:
                columns = [i for i in range(cols)]
            
            self.df = pd.DataFrame(columns=columns)
            for _ in range(rows):
                self.df.loc[len(self.df)] = ['' for _ in range(cols)]
            
            self.filtered_df = self.df.copy()
            
            self.current_file_path = None
            self.is_modified = False
            self.undo_stack.clear()
            self.redo_stack.clear()
            
            self.populate_table(self.filtered_df)
            self.update_info_panel()
            self.update_filter_combo()
            self.enable_controls(True)
            
            file_type = "Excel" if "xlsx" in settings['file_type'] else "CSV"
            self.setWindowTitle(f"Excel Editor Pro - New {file_type} File*")
            self.statusBar().showMessage(f"New {file_type} file created ({rows} rows, {cols} columns)")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error creating file: {str(e)}")
    
    def load_file(self):
        try:
            if self.is_modified:
                reply = QMessageBox.question(
                    self, 'Unsaved Changes', 
                    'You have unsaved changes. Do you want to save them first?',
                    QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel,
                    QMessageBox.Cancel
                )
                
                if reply == QMessageBox.Cancel:
                    return
                elif reply == QMessageBox.Yes:
                    if not self.save_file():
                        return
            
            file_path, _ = QFileDialog.getOpenFileName(
                self, 
                "Open Excel/CSV File", 
                "", 
                "Excel files (*.xlsx *.xls);;CSV files (*.csv);;All files (*.*)"
            )
            
            if file_path:
                self.statusBar().showMessage("Loading file...")
                self.current_file_path = file_path
                self.start_file_loading(file_path)
                self.add_to_recent_files(file_path)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error loading file: {str(e)}")
    
    def start_file_loading(self, file_path, sheet_name=None):
        try:
            self.data_processor = DataProcessor(file_path, sheet_name)
            self.data_processor.finished.connect(self.on_file_loaded)
            self.data_processor.error.connect(self.on_load_error)
            self.data_processor.progress.connect(self.statusBar().showMessage)
            self.data_processor.sheets_found.connect(self.on_sheets_found)
            self.data_processor.start()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error starting file loading: {str(e)}")
    
    def on_sheets_found(self, sheet_names):
        try:
            dialog = SheetSelectionDialog(sheet_names, self)
            if dialog.exec_() == QDialog.Accepted:
                selected_sheet = dialog.get_selected_sheet()
                self.start_file_loading(self.current_file_path, selected_sheet)
            else:
                self.statusBar().showMessage("File loading cancelled")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error handling sheets: {str(e)}")
    
    def on_file_loaded(self, df):
        try:
            self.df = df
            self.filtered_df = df.copy()
            
            if len(df) > self.MAX_DISPLAY_ROWS:
                QMessageBox.information(
                    self, 
                    "Large Dataset", 
                    f"This file has {len(df)} rows. For performance, only the first {self.MAX_DISPLAY_ROWS} rows will be displayed.\n\n"
                    "All data is still loaded and will be saved. Use filters to view specific rows."
                )
            
            self.populate_table(self.filtered_df)
            self.update_info_panel()
            self.update_filter_combo()
            self.enable_controls(True)
            self.is_modified = False
            self.undo_redo_manager.clear()
            
            file_name = os.path.basename(self.current_file_path)
            self.setWindowTitle(f"Excel Editor Pro - {file_name}")
            self.statusBar().showMessage(f"Loaded: {file_name} ({len(df)} rows, {len(df.columns)} columns)")
            self.add_to_recent_files(self.current_file_path)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error processing loaded file: {str(e)}")
    
    def on_load_error(self, error_message):
        QMessageBox.critical(self, "Error", error_message)
        self.statusBar().showMessage("Error loading file")
    
    def populate_table(self, df):
        try:
            self.table_widget.blockSignals(True)
            self.is_updating = True
            
            display_df = df.head(self.MAX_DISPLAY_ROWS) if len(df) > self.MAX_DISPLAY_ROWS else df
            
            self.table_widget.setRowCount(len(display_df))
            self.table_widget.setColumnCount(len(df.columns))
            self.table_widget.setHorizontalHeaderLabels([str(col) for col in df.columns])
            
            for i in range(len(display_df)):
                for j in range(len(df.columns)):
                    try:
                        value = display_df.iloc[i, j]
                        if pd.isna(value):
                            value = ""
                        else:
                            value = str(value)
                        
                        item = QTableWidgetItem(value)
                        
                        # Apply formatting if available for this column
                        col_name = str(df.columns[j])
                        if col_name in self.format_settings:
                            self._apply_cell_formatting(item, self.format_settings[col_name])
                        
                        self.table_widget.setItem(i, j, item)
                    except Exception:
                        item = QTableWidgetItem("")
                        self.table_widget.setItem(i, j, item)
            
            # Apply header formatting
            self._apply_header_formatting()
            
            header = self.table_widget.horizontalHeader()
            header.setSectionResizeMode(QHeaderView.Interactive)
            
            if len(display_df) <= 100:
                self.table_widget.resizeColumnsToContents()
            else:
                for col in range(len(df.columns)):
                    # Apply column width from format settings if available
                    col_name = str(df.columns[col])
                    if col_name in self.format_settings:
                        width = self.format_settings[col_name].get('width', 100)
                        if width != 'auto':
                            self.table_widget.setColumnWidth(col, width)
                        else:
                            self.table_widget.resizeColumnToContents(col)
                    else:
                        self.table_widget.setColumnWidth(col, 100)
            
            self.table_widget.blockSignals(False)
            self.is_updating = False
            
        except Exception as e:
            self.table_widget.blockSignals(False)
            self.is_updating = False
            QMessageBox.critical(self, "Error", f"Error populating table: {str(e)}")
    
    def update_table(self):
        """Update the table view with current dataframe state"""
        if self.filtered_df is not None:
            self.populate_table(self.filtered_df)
        elif self.df is not None:
            self.populate_table(self.df)

    #def on_rows_moved(self, parent, start, end, destination, row):
        #"""Handle row reordering via drag and drop and sync with DataFrame"""
        #try:
            #if self.df is None or self.is_updating:
            #    return
            
            # destination 'row' is the index where the rows are moved to
            # We need to find the actual index in the main DataFrame
            #if len(self.filtered_df) > start:
            #    source_actual_index = self.filtered_df.index[start]
                
                # Get the row to move
            #    row_to_move = self.df.loc[[source_actual_index]]
                
                # Remove from current position
            #    self.df = self.df.drop(source_actual_index)
                
                # Determine target insertion point in the main DataFrame
            #    if row >= len(self.filtered_df):
                    # Move to the very end
            #        self.df = pd.concat([self.df, row_to_move], ignore_index=True)
            #    else:
                    # Find the actual index of the row currently at the destination position
            #       target_actual_index = self.filtered_df.index[row]
                    # Find where this index is in the current (already dropped) df
            #        main_idx = self.df.index.get_loc(target_actual_index)
                    
            #        df_top = self.df.iloc[:main_idx]
            #        df_bottom = self.df.iloc[main_idx:]
            #        self.df = pd.concat([df_top, row_to_move, df_bottom], ignore_index=True)
                
            #    self.apply_filter()
            #    self.mark_as_modified()
            #    self.statusBar().showMessage("Row reordered successfully")
                
        #except Exception as e:
        #    print(f"Error syncing row move: {e}")
    
    def _apply_cell_formatting(self, item, settings):
        """Apply formatting settings to a QTableWidgetItem"""
        try:
            # Font settings
            font = QtGui.QFont()
            font.setFamily(settings.get('font_name', 'Calibri'))
            font.setPointSize(settings.get('font_size', 11))
            font.setBold(settings.get('bold', False))
            font.setItalic(settings.get('italic', False))
            font.setUnderline(settings.get('underline', False))
            item.setFont(font)
            
            # Text color
            text_color = settings.get('text_color', '#000000')
            item.setForeground(QtGui.QColor(text_color))
            
            # Background color
            bg_color = settings.get('bg_color', '#FFFFFF')
            if bg_color and bg_color != '#FFFFFF':
                item.setBackground(QtGui.QColor(bg_color))
            
            # Text alignment
            h_align = settings.get('h_align', 'general')
            v_align = settings.get('v_align', 'center')
            
            alignment = Qt.AlignVCenter  # Default vertical center
            if v_align == 'top':
                alignment = Qt.AlignTop
            elif v_align == 'bottom':
                alignment = Qt.AlignBottom
            
            if h_align == 'left':
                alignment |= Qt.AlignLeft
            elif h_align == 'center':
                alignment |= Qt.AlignHCenter
            elif h_align == 'right':
                alignment |= Qt.AlignRight
            else:  # general
                alignment |= Qt.AlignLeft
            
            item.setTextAlignment(alignment)
            
        except Exception as e:
            print(f"Error applying cell formatting: {e}")
    
    def _apply_header_formatting(self):
        """Apply header formatting to all columns that have format settings"""
        try:
            if self.df is None or self.df.empty:
                return
                
            for col_idx, col_name in enumerate(self.df.columns):
                col_name_str = str(col_name)
                if col_name_str in self.format_settings:
                    settings = self.format_settings[col_name_str]
                    
                    # Only apply header formatting if specified
                    if settings.get('format_header', True):
                        header_item = self.table_widget.horizontalHeaderItem(col_idx)
                        if not header_item:
                            header_item = QTableWidgetItem(col_name_str)
                            self.table_widget.setHorizontalHeaderItem(col_idx, header_item)
                        
                        # Header font
                        font = QtGui.QFont()
                        font.setFamily(settings.get('font_name', 'Calibri'))
                        font.setPointSize(settings.get('font_size', 11))
                        font.setBold(settings.get('header_bold', True))
                        header_item.setFont(font)
                        
                        # Header text color
                        header_text_color = settings.get('header_text_color', '#FFFFFF')
                        header_item.setForeground(QtGui.QColor(header_text_color))
                        
                        # Header background color
                        header_bg = settings.get('header_bg_color', '#4472C4')
                        header_item.setBackground(QtGui.QColor(header_bg))
                        
        except Exception as e:
            print(f"Error applying header formatting: {e}")


    
    def mark_as_modified(self):
        if not self.is_modified:
            self.is_modified = True
            current_title = self.windowTitle()
            if not current_title.endswith('*'):
                self.setWindowTitle(current_title + '*')
    
    def auto_save(self):
        try:
            if self.df is not None and self.is_modified and self.current_file_path:
                self.save_file()
                self.statusBar().showMessage("Auto-saved", 2000)
        except Exception as e:
            print(f"Auto-save error: {e}")
    
    def save_file(self):
        try:
            if self.df is None:
                return False
            
            if self.current_file_path is None:
                return self.save_as_file()
            
            include_index = self.include_index_cb.isChecked()
            
            self.statusBar().showMessage("Saving file...")
            QApplication.processEvents()
            
            if self.current_file_path.endswith('.csv'):
                self.df.to_csv(self.current_file_path, index=include_index)
            else:
                # Save with formatting if settings exist
                if self.format_settings and OPENPYXL_AVAILABLE:
                    # First save without formatting
                    self.df.to_excel(self.current_file_path, index=include_index, engine='openpyxl')
                    
                    # Then load and apply formatting
                    wb = load_workbook(self.current_file_path)
                    ws = wb.active
                    FormatApplier.apply_formatting(ws, self.format_settings, self.df)
                    FormatApplier.apply_borders(ws, self.df)
                    wb.save(self.current_file_path)
                else:
                    self.df.to_excel(self.current_file_path, index=include_index, engine='openpyxl')
            
            self.is_modified = False
            file_name = os.path.basename(self.current_file_path)
            self.setWindowTitle(f"Excel Editor Pro- {file_name}")
            self.statusBar().showMessage(f"File saved: {file_name}")
            return True
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error saving file: {str(e)}")
            return False
    
    def save_as_file(self):
        try:
            if self.df is None:
                return False
            
            file_path, file_type = QFileDialog.getSaveFileName(
                self,
                "Save File As",
                "",
                "Excel files (*.xlsx);;CSV files (*.csv);;All files (*.*)"
            )
            
            if file_path:
                include_index = self.include_index_cb.isChecked()
                
                self.statusBar().showMessage("Saving file...")
                QApplication.processEvents()
                
                if file_path.endswith('.csv'):
                    self.df.to_csv(file_path, index=include_index)
                else:
                    if not file_path.endswith('.xlsx'):
                        file_path += '.xlsx'
                    # Save with formatting if settings exist
                    if self.format_settings and OPENPYXL_AVAILABLE:
                        # First save without formatting
                        self.df.to_excel(file_path, index=include_index, engine='openpyxl')
                        
                        # Then load and apply formatting
                        wb = load_workbook(file_path)
                        ws = wb.active
                        FormatApplier.apply_formatting(ws, self.format_settings, self.df)
                        FormatApplier.apply_borders(ws, self.df)
                        wb.save(file_path)
                    else:
                        self.df.to_excel(file_path, index=include_index, engine='openpyxl')
                
                self.current_file_path = file_path
                self.is_modified = False
                file_name = os.path.basename(self.current_file_path)
                self.setWindowTitle(f"Excel Editor Pro - {file_name}")
                self.statusBar().showMessage(f"File saved: {file_name}")
                self.add_to_recent_files(self.current_file_path)
                QMessageBox.information(self, "Success", "File saved successfully!")
                return True
                
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error saving file: {str(e)}")
            return False
        return False

    def export_to_doc(self):
        try:
            if self.df is None or self.df.empty:
                QMessageBox.warning(self, "Warning", "No data to export.")
                return

            # Default required columns for pre-selection
            default_cols = ["full name", "gender", "phone", "member type"]
            
            # Show column selection dialog
            dialog = ColumnSelectionDialog(self.df.columns, default_selected=default_cols, parent=self)
            if dialog.exec_() != QDialog.Accepted:
                return
                
            selected_cols = dialog.get_selected_columns()
            if not selected_cols:
                QMessageBox.warning(self, "Warning", "No columns selected for export.")
                return

            file_path, _ = QFileDialog.getSaveFileName(
                self, "Export Report as DOC", "", "Word Documents (*.docx)"
            )
            
            if not file_path:
                return

            if not file_path.endswith('.docx'):
                file_path += '.docx'

            self.statusBar().showMessage("Exporting to DOC...")
            QApplication.processEvents()

            doc = Document()
            
            # Add title
            title = doc.add_heading('Member Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add table (with an extra column for Index)
            table = doc.add_table(rows=1, cols=len(selected_cols) + 1)
            table.style = 'Table Grid'
            
            # Set headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "No."
            for i, col_name in enumerate(selected_cols):
                hdr_cells[i + 1].text = str(col_name).title()
            
            # Make all headers bold
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            # Add data rows
            for idx, row in enumerate(self.df.iterrows(), 1):
                _, row_data = row
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx) # Index column
                for i, col_name in enumerate(selected_cols):
                    val = row_data[col_name]
                    row_cells[i + 1].text = str(val) if not pd.isna(val) else ""

            doc.save(file_path)
            
            self.statusBar().showMessage(f"Report exported: {os.path.basename(file_path)}")
            QMessageBox.information(self, "Success", f"Report exported successfully to {os.path.basename(file_path)}")
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error exporting report: {str(e)}")
    
    def update_info_panel(self):
        try:
            if self.df is not None:
                info = []
                info.append(f"Total Rows: {len(self.df)}")
                info.append(f"Total Columns: {len(self.df.columns)}")
                info.append(f"Displayed Rows: {min(len(self.filtered_df), self.MAX_DISPLAY_ROWS)}")
                
                try:
                    mem_usage = self.df.memory_usage(deep=True).sum() / (1024 * 1024)
                    if mem_usage < 1:
                        info.append(f"Memory Usage: {mem_usage * 1024:.1f} KB")
                    else:
                        info.append(f"Memory Usage: {mem_usage:.1f} MB")
                except:
                    pass
                
                dtype_counts = self.df.dtypes.value_counts()
                info.append("\nColumn Types:")
                for dtype, count in dtype_counts.items():
                    info.append(f"  {dtype}: {count}")
                
                try:
                    missing_values = self.df.isnull().sum()
                    missing_total = missing_values.sum()
                    if missing_total > 0:
                        info.append(f"\nMissing Values: {missing_total}")
                except:
                    pass
                
                self.info_text.setPlainText("\n".join(info))
        except Exception as e:
            self.info_text.setPlainText(f"Error updating info: {str(e)}")
    
    def update_filter_combo(self):
        try:
            if self.df is not None:
                self.filter_column_combo.clear()
                self.filter_column_combo.addItems([str(col) for col in self.df.columns])
        except Exception as e:
            print(f"Error updating filter combo: {e}")
    
    def apply_filter(self):
        try:
            if self.df is None:
                return
            
            filter_text = self.filter_input.text().strip()
            if not filter_text:
                self.filtered_df = self.df.copy()
            else:
                column = self.filter_column_combo.currentText()
                if column in self.df.columns:
                    mask = self.df[column].astype(str).str.contains(filter_text, case=False, na=False)
                    self.filtered_df = self.df[mask].copy()
                else:
                    self.filtered_df = self.df.copy()
            
            self.populate_table(self.filtered_df)
            self.update_info_panel()
            self.clear_filter_btn.setEnabled(bool(filter_text))
            
            if len(self.filtered_df) < len(self.df):
                self.statusBar().showMessage(f"Filter applied: {len(self.filtered_df)} of {len(self.df)} rows shown")
            else:
                self.statusBar().showMessage("Filter cleared - showing all rows")
                
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error applying filter: {str(e)}")
    
    def clear_filter(self):
        self.filter_input.clear()
        self.apply_filter()
    
    def add_row(self):
        try:
            if self.df is None:
                return
            
            dialog = AddRowDialog(self.df.columns.tolist(), self)
            if dialog.exec_() == QDialog.Accepted:
                # Save state for undo
                self.undo_redo_manager.save_state("Add Row")
                
                new_data = dialog.get_data()
                new_row = pd.DataFrame([new_data])
                self.df = pd.concat([self.df, new_row], ignore_index=True)
                self.apply_filter()
                self.mark_as_modified()
                self.statusBar().showMessage("Row added successfully")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error adding row: {str(e)}")
    
    def add_column(self):
        try:
            if self.df is None:
                return
            
            dialog = AddColumnDialog(self.df.columns.tolist(), self)
            if dialog.exec_() == QDialog.Accepted:
                data = dialog.get_data()
                column_name = data['name']
                position = data['position']
                default_value = data['default_value'] if data['default_value'] else None
                
                if not column_name:
                    QMessageBox.warning(self, "Invalid Input", "Column name cannot be empty!")
                    return
                
                if column_name in self.df.columns:
                    QMessageBox.warning(self, "Duplicate Column", "Column name already exists!")
                    return
                
                # Save state for undo
                self.undo_redo_manager.save_state("Add Column")
                
                if position == 0:
                    self.df[column_name] = default_value
                else:
                    insert_loc = position - 1
                    cols = list(self.df.columns)
                    cols.insert(insert_loc, column_name)
                    self.df[column_name] = default_value
                    self.df = self.df[cols]
                
                self.apply_filter()
                self.update_filter_combo()
                self.mark_as_modified()
                self.statusBar().showMessage(f"Column '{column_name}' added successfully")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error adding column: {str(e)}")
    
    def format_columns(self):
        """Open column formatting dialog"""
        try:
            if self.df is None or self.df.empty:
                QMessageBox.warning(self, "No Data", "Please load a file first.")
                return
            
            dialog = ColumnFormattingDialog(self.df.columns.tolist(), self)
            if dialog.exec_() == QDialog.Accepted:
                self.format_settings = dialog.get_format_settings()
                # Refresh the table to show the formatting changes immediately
                self.populate_table(self.filtered_df)
                self.mark_as_modified()
                self.statusBar().showMessage("Column formatting applied successfully!")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error opening format dialog: {str(e)}")
    
    def open_advanced_formatting(self):
        """Open advanced formatting dialog with multiple features"""
        try:
            if self.df is None or self.df.empty:
                QMessageBox.warning(self, "No Data", "Please load a file first.")
                return
            
            dialog = AdvancedFormattingDialog(self, self.table_widget, self.df)
            dialog.exec_()
            
            # Mark as modified if any changes were made
            self.mark_as_modified()
            self.statusBar().showMessage("Advanced formatting dialog closed")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error opening advanced formatting: {str(e)}")
    
    def open_data_transformation(self):
        """Open data transformation dialog"""
        try:
            if self.df is None or self.df.empty:
                QMessageBox.warning(self, "No Data", "Please load a file first.")
                return
            
            dialog = DataTransformationDialog(self.df, self)
            if dialog.exec_() == QDialog.Accepted:
                result_df = dialog.get_result()
                if result_df is not None:
                    self.df = result_df
                    self.filtered_df = self.df.copy()
                    self.populate_table(self.filtered_df)
                    self.update_info_panel()
                    self.update_filter_combo()
                    self.mark_as_modified()
                    self.statusBar().showMessage("Data transformation applied successfully")
                    QMessageBox.information(self, "Success", 
                                          f"Transformation complete!\nNew shape: {len(self.df)} rows × {len(self.df.columns)} columns")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error in data transformation: {str(e)}\n\n{traceback.format_exc()}")
    
    def delete_row(self):
        try:
            if self.df is None:
                return
            
            current_row = self.table_widget.currentRow()
            if current_row >= 0:
                reply = QMessageBox.question(
                    self, 
                    'Delete Row', 
                    'Are you sure you want to delete the selected row?',
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.No
                )
                
                if reply == QMessageBox.Yes:
                    # Save state for undo
                    self.undo_redo_manager.save_state("Delete Row")
                    
                    if len(self.filtered_df) > current_row:
                        actual_index = self.filtered_df.index[current_row]
                        self.df = self.df.drop(actual_index).reset_index(drop=True)
                        self.apply_filter()
                        self.mark_as_modified()
                        self.statusBar().showMessage("Row deleted successfully")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error deleting row: {str(e)}")

    def move_row_up(self):
        """Move selected row up by one position"""
        try:
            if self.df is None:
                return
            
            current_row = self.table_widget.currentRow()
            if current_row <= 0:  # Already at top or no selection
                self.statusBar().showMessage("Row is already at the top")
                return
            
            if len(self.filtered_df) > current_row:
                # Get current and previous row indices from filtered dataframe
                current_actual_index = self.filtered_df.index[current_row]
                previous_actual_index = self.filtered_df.index[current_row - 1]
                
                # Get their positions in the main dataframe
                current_pos = self.df.index.get_loc(current_actual_index)
                previous_pos = self.df.index.get_loc(previous_actual_index)
                
                # Get the rows
                row_current = self.df.loc[[current_actual_index]].copy()
                row_previous = self.df.loc[[previous_actual_index]].copy()
                
                # Drop both rows
                self.df = self.df.drop([current_actual_index, previous_actual_index])
                
                # Determine insertion point (the earlier position)
                insert_pos = min(current_pos, previous_pos)
                
                # Split dataframe and reinsert in swapped order
                df_before = self.df.iloc[:insert_pos]
                df_after = self.df.iloc[insert_pos:]
                
                # Insert current row before previous row
                self.df = pd.concat([df_before, row_current, row_previous, df_after], ignore_index=True)
                
                self.apply_filter()
                self.mark_as_modified()
                self.table_widget.selectRow(current_row - 1)
                self.statusBar().showMessage("Row moved up")
                
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error moving row up: {str(e)}")


    def move_row_down(self):
        """Move selected row down by one position"""
        try:
            if self.df is None:
                return
            
            current_row = self.table_widget.currentRow()
            if current_row < 0 or current_row >= len(self.filtered_df) - 1:
                self.statusBar().showMessage("Row is already at the bottom")
                return
            
            if len(self.filtered_df) > current_row + 1:
                # Get current and next row indices from filtered dataframe
                current_actual_index = self.filtered_df.index[current_row]
                next_actual_index = self.filtered_df.index[current_row + 1]
                
                # Get their positions in the main dataframe
                current_pos = self.df.index.get_loc(current_actual_index)
                next_pos = self.df.index.get_loc(next_actual_index)
                
                # Get the rows
                row_current = self.df.loc[[current_actual_index]].copy()
                row_next = self.df.loc[[next_actual_index]].copy()
                
                # Drop both rows
                self.df = self.df.drop([current_actual_index, next_actual_index])
                
                # Determine insertion point (the earlier position)
                insert_pos = min(current_pos, next_pos)
                
                # Split dataframe and reinsert in swapped order
                df_before = self.df.iloc[:insert_pos]
                df_after = self.df.iloc[insert_pos:]
                
                # Insert next row before current row
                self.df = pd.concat([df_before, row_next, row_current, df_after], ignore_index=True)
                
                self.apply_filter()
                self.mark_as_modified()
                self.table_widget.selectRow(current_row + 1)
                self.statusBar().showMessage("Row moved down")
                
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error moving row down: {str(e)}")
    
    def show_statistics(self):
        try:
            if self.df is not None:
                dialog = StatisticsDialog(self.df, self)
                dialog.exec_()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error showing statistics: {str(e)}")
    
    def open_settings(self):
        """Open settings dialog for theme and auto-save"""
        try:
            dialog = SettingsDialog(
                self.bg_color, 
                self.accent_color, 
                self.auto_save_enabled, 
                self.auto_save_interval,
                self.auto_save_manager.keep_versions,
                self
            )
            if dialog.exec_() == QDialog.Accepted:
                settings = dialog.get_settings()
                
                # Update theme
                bg, accent = settings['bg_color'], settings['accent_color']
                self.bg_color = bg
                self.accent_color = accent
                self.settings.setValue("theme/bg_color", bg)
                self.settings.setValue("theme/accent_color", accent)
                apply_custom_theme(QApplication.instance(), bg, accent)
                
                # Update auto-save
                self.auto_save_enabled = settings['auto_save_enabled']
                self.auto_save_interval = settings['auto_save_interval']
                self.settings.setValue("auto_save/enabled", self.auto_save_enabled)
                self.settings.setValue("auto_save/interval", self.auto_save_interval)
                
                # Update auto-save manager with all settings including keep_versions
                self.auto_save_manager.update_settings(
                    settings['auto_save_enabled'],
                    settings['auto_save_interval'],
                    settings['keep_versions']
                )
                
                QMessageBox.information(
                    self,
                    "Settings Updated",
                    "Settings have been updated successfully!"
                )
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error updating settings: {str(e)}")
    
    def enable_controls(self, enabled):
        try:
            self.save_btn.setEnabled(enabled)
            self.save_as_btn.setEnabled(enabled)
            self.add_row_btn.setEnabled(enabled)
            self.add_column_btn.setEnabled(enabled)
            self.format_column_btn.setEnabled(enabled)
            self.advanced_format_btn.setEnabled(enabled)
            self.transform_btn.setEnabled(enabled)
            self.visualize_btn.setEnabled(enabled)
            self.ai_btn.setEnabled(enabled)
            self.delete_row_btn.setEnabled(enabled)
            self.move_up_btn.setEnabled(enabled)  # Changed
            self.move_down_btn.setEnabled(enabled)  # Changed
            self.stats_btn.setEnabled(enabled)
            self.apply_filter_btn.setEnabled(enabled)
            self.clear_filter_btn.setEnabled(enabled and bool(self.filter_input.text()))
        except Exception as e:
            print(f"Error enabling controls: {e}")
    
    def refresh_view(self):
        """Refresh the table view"""
        try:
            if self.df is not None:
                self.populate_table(self.filtered_df)
                self.statusBar().showMessage("View refreshed")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error refreshing view: {str(e)}")
    
    def show_help(self):
        """Show the help dialog"""
        try:
            help_dialog = HelpDialog(self)
            help_dialog.show()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error opening help: {str(e)}")
    
    def show_keyboard_shortcuts(self):
        """Show keyboard shortcuts in help dialog"""
        try:
            help_dialog = HelpDialog(self)
            help_dialog.tab_widget.setCurrentIndex(1)  # Switch to keyboard shortcuts tab
            help_dialog.show()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error opening keyboard shortcuts: {str(e)}")
    
    def show_about(self):
        """Show about dialog"""
        try:
            help_dialog = HelpDialog(self)
            help_dialog.tab_widget.setCurrentIndex(6)  # Switch to about tab
            help_dialog.show()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error opening about: {str(e)}")
    
    def setup_keyboard_shortcuts(self):
        """Setup global keyboard shortcuts"""
        try:
            # Filter focus
            filter_shortcut = QShortcut(QKeySequence('Ctrl+F'), self)
            filter_shortcut.activated.connect(lambda: self.filter_input.setFocus())
            
            # Help
            help_shortcut = QShortcut(QKeySequence('F1'), self)
            help_shortcut.activated.connect(self.show_help)
            
            # Refresh
            refresh_shortcut = QShortcut(QKeySequence('F5'), self)
            refresh_shortcut.activated.connect(self.refresh_view)
            
        except Exception as e:
            print(f"Error setting up keyboard shortcuts: {e}")
    
    def closeEvent(self, event):
        try:
            # Close split view if open
            if hasattr(self, 'split_view_manager'):
                self.split_view_manager.close_split_view()
            
            if self.is_modified:
                reply = QMessageBox.question(
                    self, 'Unsaved Changes', 
                    'You have unsaved changes. Do you want to save them before closing?',
                    QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel,
                    QMessageBox.Cancel
                )
                
                if reply == QMessageBox.Cancel:
                    event.ignore()
                    return
                elif reply == QMessageBox.Yes:
                    if not self.save_file():
                        event.ignore()
                        return
            
            event.accept()
        except Exception as e:
            print(f"Error in close event: {e}")
            event.accept()
    
    # ========== NEW PRODUCTIVITY FEATURES ==========
    
    def undo(self):
        """Undo last action"""
        try:
            self.undo_redo_manager.undo()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error undoing: {str(e)}")
    
    def redo(self):
        """Redo last undone action"""
        try:
            self.undo_redo_manager.redo()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error redoing: {str(e)}")
    
    def open_edit_history(self):
        """Open edit history dialog"""
        try:
            history_dialog = UndoRedoHistoryDialog(self, self.undo_redo_manager)
            history_dialog.exec_()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error opening edit history: {str(e)}")
    
    def open_version_history(self):
        """Open version history dialog"""
        try:
            if not self.current_file_path:
                QMessageBox.information(
                    self,
                    "No File",
                    "Please open a file first to view its version history."
                )
                return
            
            version_dialog = VersionHistoryDialog(self, self.auto_save_manager)
            version_dialog.exec_()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error opening version history: {str(e)}")
    
    def open_autosave_settings(self):
        """Open auto-save settings dialog"""
        try:
            settings_dialog = AutoSaveSettingsDialog(self, self.auto_save_manager)
            if settings_dialog.exec_():
                self.statusBar().showMessage("Auto-save settings updated", 2000)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error opening auto-save settings: {str(e)}")
    
    def open_favorites(self):
        """Open favorites dialog"""
        try:
            favorites_dialog = FavoritesDialog(self, self.favorites_manager)
            favorites_dialog.exec_()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error opening favorites: {str(e)}")
    
    def add_current_to_favorites(self):
        """Add current file to favorites"""
        try:
            if not self.current_file_path:
                QMessageBox.information(
                    self,
                    "No File",
                    "No file is currently open."
                )
                return
            
            self.favorites_manager.add_favorite(self.current_file_path)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error adding to favorites: {str(e)}")
    
    def open_freeze_panes(self):
        """Open freeze panes dialog"""
        try:
            if self.df is None:
                QMessageBox.information(
                    self,
                    "No Data",
                    "Please load a file first to use freeze panes."
                )
                return
            
            freeze_dialog = ColumnFreezeDialog(self, self.freeze_manager)
            freeze_dialog.exec_()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error opening freeze panes: {str(e)}")
    
    # ========== AI FEATURES ==========
    
    def open_ai_features(self):
        """Open AI features dialog"""
        try:
            if self.df is None:
                QMessageBox.information(
                    self,
                    "No Data",
                    "Please load a file first to use AI features."
                )
                return
            
            ai_dialog = AIFeaturesDialog(self, self.df)
            ai_dialog.exec_()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error opening AI features: {str(e)}")

    # Visualization Methods
    def open_visualization_menu(self):
        """Open a context menu for visualization options"""
        menu = QMenu(self)
        
        chart_action = QAction("📊 Create Chart", self)
        chart_action.triggered.connect(self.create_chart)
        menu.addAction(chart_action)
        
        dashboard_action = QAction("🚀 Dashboard View", self)
        dashboard_action.triggered.connect(self.open_dashboard)
        menu.addAction(dashboard_action)
        
        # Show menu under the button
        menu.exec_(self.visualize_btn.mapToGlobal(self.visualize_btn.rect().bottomLeft()))

    def create_chart(self):
        """Open chart creation dialog"""
        if not VisualizationManager.is_available():
            QMessageBox.warning(self, "Missing Dependencies", VisualizationManager.get_missing_deps_message())
            return
            
        if self.df is None or self.df.empty:
            QMessageBox.warning(self, "No Data", "Please load a file with data first.")
            return
            
        try:
            dialog = ChartDialog(self.df, self)
            dialog.exec_()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error opening chart dialog: {str(e)}")

    def open_dashboard(self):
        """Open dashboard view"""
        if not VisualizationManager.is_available():
            QMessageBox.warning(self, "Missing Dependencies", VisualizationManager.get_missing_deps_message())
            return
            
        if self.df is None or self.df.empty:
            QMessageBox.warning(self, "No Data", "Please load a file with data first.")
            return
            
        try:
            dialog = DashboardDialog(self.df, self)
            dialog.exec_()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error opening dashboard: {str(e)}")
    
    def quick_ai_analysis(self):
        """Quick AI analysis and show summary"""
        try:
            if self.df is None:
                QMessageBox.information(self, "No Data", "Please load a file first.")
                return
            
            from AIFeatures_ import AIManager, SKLEARN_AVAILABLE
            
            if not SKLEARN_AVAILABLE:
                QMessageBox.warning(
                    self, 
                    "AI Not Available",
                    "AI features require additional packages.\n\n"
                    "Install with: pip install scikit-learn scipy"
                )
                return
            
            # Show progress
            self.statusBar().showMessage("Running AI analysis...")
            QApplication.processEvents()
            
            # Run analysis
            ai_manager = AIManager(self.df)
            results = ai_manager.get_full_analysis()
            
            # Show summary
            summary = "🤖 AI Analysis Summary\n\n"
            
            if 'data_quality' in results:
                quality = results['data_quality']
                summary += f"Data Quality Score: {quality['score']:.1f}/100\n\n"
            
            if 'insights' in results:
                summary += "Key Insights:\n"
                for insight in results['insights'][:5]:
                    summary += f"• {insight}\n"
                summary += "\n"
            
            if 'column_types' in results:
                summary += f"Detected {len(results['column_types'])} column types\n"
            
            self.statusBar().showMessage("AI analysis complete", 3000)
            QMessageBox.information(self, "AI Analysis", summary)
            
        except Exception as e:
            self.statusBar().clearMessage()
            QMessageBox.critical(self, "Error", f"AI analysis failed: {str(e)}")
    
    def smart_clean_data(self):
        """Smart data cleaning with AI suggestions"""
        try:
            if self.df is None:
                QMessageBox.information(self, "No Data", "Please load a file first.")
                return
            
            from AIFeatures_ import AIManager, SKLEARN_AVAILABLE
            
            if not SKLEARN_AVAILABLE:
                QMessageBox.warning(
                    self, 
                    "AI Not Available",
                    "This feature requires: pip install scikit-learn scipy"
                )
                return
            
            # Open AI dialog directly to cleaning tab
            ai_dialog = AIFeaturesDialog(self, self.df)
            ai_dialog.tabs.setCurrentIndex(2)  # Cleaning tab
            ai_dialog.exec_()
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error opening smart clean: {str(e)}")
    
    def get_ai_suggestions(self):
        """Get AI suggestions for current data"""
        try:
            if self.df is None:
                QMessageBox.information(self, "No Data", "Please load a file first.")
                return
            
            from AIFeatures_ import AIManager, SKLEARN_AVAILABLE
            
            if not SKLEARN_AVAILABLE:
                QMessageBox.warning(
                    self, 
                    "AI Not Available",
                    "This feature requires: pip install scikit-learn scipy"
                )
                return
            
            # Get AI suggestions
            self.statusBar().showMessage("Generating AI suggestions...")
            QApplication.processEvents()
            
            ai_manager = AIManager(self.df)
            
            # Get general suggestions
            suggestions_text = "💡 AI Suggestions for Your Data\n\n"
            
            # Check for missing data
            missing_cols = [col for col in self.df.columns if self.df[col].isna().sum() > 0]
            if missing_cols:
                suggestions_text += "Missing Data:\n"
                for col in missing_cols[:5]:
                    missing_pct = self.df[col].isna().sum() / len(self.df) * 100
                    suggestions_text += f"• {col}: {missing_pct:.1f}% missing - consider filling\n"
                suggestions_text += "\n"
            
            # Check for numeric columns
            numeric_cols = self.df.select_dtypes(include=['number']).columns.tolist()
            if len(numeric_cols) >= 2:
                suggestions_text += "Calculations:\n"
                suggestions_text += f"• Try combining {numeric_cols[0]} and {numeric_cols[1]}\n"
                suggestions_text += "• Use AI predictions for missing values\n\n"
            
            # General suggestions
            suggestions_text += "Actions:\n"
            suggestions_text += "• Run full AI analysis for detailed insights\n"
            suggestions_text += "• Use smart cleaning for data quality issues\n"
            suggestions_text += "• Try AI-powered predictions\n"
            
            self.statusBar().clearMessage()
            QMessageBox.information(self, "AI Suggestions", suggestions_text)
            
        except Exception as e:
            self.statusBar().clearMessage()
            QMessageBox.critical(self, "Error", f"Error getting suggestions: {str(e)}")
    
    # ========== CLOUD SYNC FEATURES ==========
    
    def open_cloud_sync(self):
        """Open cloud synchronization dialog"""
        try:
            cloud_dialog = CloudSyncDialog(self)
            cloud_dialog.exec_()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error opening cloud sync: {str(e)}")
    
    def quick_upload_to_cloud(self):
        """Quick upload current file to cloud"""
        try:
            if not self.current_file_path:
                QMessageBox.information(
                    self,
                    "No File",
                    "Please save your file first before uploading to cloud."
                )
                return
            
            # Open cloud sync dialog directly to upload tab
            cloud_dialog = CloudSyncDialog(self)
            cloud_dialog.use_current_file()
            cloud_dialog.tabs.setCurrentIndex(1)  # Upload tab
            cloud_dialog.exec_()
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error uploading to cloud: {str(e)}")
    
    def download_from_cloud(self):
        """Download file from cloud storage"""
        try:
            # Open cloud sync dialog directly to download tab
            cloud_dialog = CloudSyncDialog(self)
            cloud_dialog.tabs.setCurrentIndex(2)  # Download tab
            cloud_dialog.exec_()
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error downloading from cloud: {str(e)}")
    
    def configure_cloud_backup(self):
        """Configure automatic cloud backup settings"""
        try:
            # Open cloud sync dialog directly to auto-sync tab
            cloud_dialog = CloudSyncDialog(self)
            cloud_dialog.tabs.setCurrentIndex(3)  # Auto-sync tab
            cloud_dialog.exec_()
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error configuring cloud backup: {str(e)}")

    # ========== DATA VALIDATION ==========

    def open_validation_dialog(self):
        """Open dialog to set validation rules for the current column"""
        try:
            if self.df is None: return
            
            current_col = self.table_widget.currentColumn()
            if current_col < 0:
                QMessageBox.warning(self, "No Selection", "Please select a cell/column first.")
                return

            dialog = ValidationDialog(self)
            if dialog.exec_() == QDialog.Accepted:
                rule = dialog.get_rule()
                if rule:
                    # Save rule
                    self.validation_manager.add_rule(current_col, rule)
                    
                    # If list type, set delegate
                    if rule.type == ValidationRule.TYPE_LIST:
                        options = rule.params.get('options', [])
                        delegate = DropdownDelegate(self.table_widget, options)
                        self.table_widget.setItemDelegateForColumn(current_col, delegate)
                    else:
                        # Remove delegate if it exists? Or keep default? 
                        # For now, we only set specific delegates. 
                        # To remove, we'd set to None or default, but PyQt doesn't easily "unset" without creating a new default one.
                        # Ideally we track delegates. 
                        self.table_widget.setItemDelegateForColumn(current_col, None) # Reset to default
                    
                    self.statusBar().showMessage(f"Validation rule applied to column {current_col + 1}")
                    
                    # Re-validate existing data
                    self.validate_data()

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error setting validation: {str(e)}")

    def validate_data(self):
        """Check all data against rules"""
        if self.df is None: return
        
        try:
            for row in range(self.table_widget.rowCount()):
                for col in range(self.table_widget.columnCount()):
                    item = self.table_widget.item(row, col)
                    if not item: continue
                    
                    value = item.text()
                    is_valid, msg = self.validation_manager.validate_cell(col, value)
                    
                    if not is_valid:
                        item.setBackground(QColor("#FFCDD2")) # Light Red
                        item.setToolTip(f"Invalid: {msg}")
                    else:
                        # Restore background (need to check if it had other formatting)
                        # For now, just clear the "error" red if it was red, or reset to default/theme
                        # This is a bit tricky if we have other colors. 
                        # Simple approach: Check if it is the error color, if so, clear it.
                        if item.background().color() == QColor("#FFCDD2"):
                            item.setBackground(QBrush(Qt.NoBrush))
                            item.setToolTip("")
        except Exception as e:
            print(f"Validation error: {e}")

    def on_cell_changed(self, row, column):
        """Handle cell changes to validate content and update data"""
        try:
            if self.df is None or self.is_updating:
                return
            
            item = self.table_widget.item(row, column)
            if not item: return
            
            new_value = item.text()
            
            # --- Update DataFrame Logic ---
            if len(self.filtered_df) > row and len(self.df.columns) > column:
                actual_index = self.filtered_df.index[row] if len(self.filtered_df) > row else row
                col_name = self.df.columns[column]
                
                if new_value == "":
                    converted_value = None
                else:
                    try:
                        if '.' in new_value:
                            converted_value = float(new_value)
                        else:
                            try:
                                converted_value = int(new_value)
                            except ValueError:
                                converted_value = new_value
                    except (ValueError, TypeError):
                        converted_value = new_value
                
                self.df.loc[actual_index, col_name] = converted_value
                self.mark_as_modified()
            
            # --- Validation Logic ---
            is_valid, msg = self.validation_manager.validate_cell(column, new_value)
            
            if not is_valid:
                item.setBackground(QColor("#FFCDD2")) # Light Red
                item.setToolTip(f"Invalid: {msg}")
                self.statusBar().showMessage(f"⚠️ Validation Error: {msg}")
            else:
                 # Clear validation error style if it was set
                 if item.background().color() == QColor("#FFCDD2"):
                    item.setBackground(QBrush(Qt.NoBrush))
                    item.setToolTip("")
                    self.statusBar().showMessage("Valid")
                    
        except Exception as e:
            print(f"Error acting on cell change: {e}")

